import os
import json
import sqlite3
import time
from io import BytesIO
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file, Response, stream_with_context
from flask_session import Session
from dotenv import load_dotenv
from google_auth_oauthlib.flow import Flow
from google.oauth2 import id_token, credentials as google_credentials
from google.auth.transport import requests as google_requests
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import pathlib
import requests
from notion_client import Client as NotionClient
from elevenlabs.client import ElevenLabs

# ── Environment ───────────────────────────────────────────────────────────────

load_dotenv()

ENVIRONMENT = os.environ.get("ENVIRONMENT", "development").lower()
IS_PRODUCTION = ENVIRONMENT == "production"

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key-change-in-prod")

# ── Session Backend ──────────────────────────────────────────────────────────
# Filesystem sessions are fine for single-server deployments (Railway, Fly.io,
# Heroku). Swap to Redis if you scale to multiple dynos/instances.

app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
if IS_PRODUCTION:
    app.config["SESSION_COOKIE_SECURE"] = True
Session(app)

# ── Error tracking ───────────────────────────────────────────────────────────
# Optional. Set SENTRY_DSN in prod env if you want error monitoring.
SENTRY_DSN = os.environ.get("SENTRY_DSN", "")
if IS_PRODUCTION and SENTRY_DSN:
    import sentry_sdk
    from sentry_sdk.integrations.flask import FlaskIntegration
    sentry_sdk.init(
        dsn=SENTRY_DSN,
        integrations=[FlaskIntegration()],
        environment=ENVIRONMENT,
        traces_sample_rate=0.1,
    )

# ── Rate Limiting ───────────────────────────────────────────────────────────

# We install a per-request rate limiter function used in routes
# (Flask-Limiter extension imported inline to avoid breaking)

# ── Security Headers (manual middleware) ────────────────────────────────────

@app.after_request
def add_security_headers(response):
    """Apply security headers to every response."""
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("X-XSS-Protection", "1; mode=block")
    response.headers.setdefault(
        "Content-Security-Policy",
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://apis.google.com; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
        "font-src 'self' https://fonts.gstatic.com; "
        "img-src 'self' https: data:; "
        "connect-src 'self' http://localhost:3000 https://*.googleapis.com; "
        "frame-src https://accounts.google.com; "
        "object-src 'none'"
    )
    if IS_PRODUCTION:
        response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
        response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    return response

# ── DeepSeek AI ──────────────────────────────────────────────────────────────

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_URL     = "https://api.deepseek.com/chat/completions"
DEEPSEEK_MODEL   = os.environ.get("DEEPSEEK_MODEL", "deepseek-v4-flash")

# Speech work is creative rather than analytical, so thinking tokens mostly buy
# latency here. Raise with DEEPSEEK_REASONING_EFFORT ('low' | 'high' | 'max').
DEEPSEEK_EFFORT = os.environ.get("DEEPSEEK_REASONING_EFFORT", "low")

# Kept under gunicorn's 120s worker timeout (see Procfile) so a slow upstream
# returns an error instead of taking the worker down with it.
DEEPSEEK_TIMEOUT = 100

DEEPSEEK_SYSTEM = (
    "You are a World Class Public Speaking Coach and Speechwriter. "
    "Return the finished work only — no preamble, no commentary on your own "
    "work, and no markdown code fences around the output."
)

# DeepSeek honours response_format only when the prompt also mentions json.
DEEPSEEK_SYSTEM_JSON = (
    DEEPSEEK_SYSTEM + " Respond with a single valid json object and nothing else."
)


class DeepSeekError(RuntimeError):
    """DeepSeek was unreachable, unconfigured, or returned nothing usable."""


def _deepseek_chat(prompt, json_mode=False, max_tokens=8192):
    """Call DeepSeek and return the assistant's text.

    Reads `content` only. With thinking mode on, the model's reasoning arrives
    in a sibling `reasoning_content` field that must never reach the user.
    """
    if not DEEPSEEK_API_KEY:
        raise DeepSeekError("DEEPSEEK_API_KEY is not configured")

    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {"role": "system",
             "content": DEEPSEEK_SYSTEM_JSON if json_mode else DEEPSEEK_SYSTEM},
            {"role": "user", "content": prompt},
        ],
        "reasoning_effort": DEEPSEEK_EFFORT,
        "max_tokens": max_tokens,
        "temperature": 0.9,
        "stream": False,
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}

    try:
        resp = requests.post(
            DEEPSEEK_URL,
            json=payload,
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            timeout=DEEPSEEK_TIMEOUT,
        )
    except requests.RequestException as e:
        raise DeepSeekError(f"DeepSeek request failed: {e}")

    if not resp.ok:
        raise DeepSeekError(
            f"DeepSeek API error {resp.status_code}: {resp.text[:300]}"
        )

    choice = (resp.json().get("choices") or [{}])[0]
    text = ((choice.get("message") or {}).get("content") or "").strip()

    if not text:
        # JSON mode can return empty content; finish_reason separates a
        # truncation from a refusal.
        raise DeepSeekError(
            f"DeepSeek returned no text (finish_reason: {choice.get('finish_reason')})"
        )
    return text


def _deepseek_json(prompt, max_tokens=8192):
    """Call DeepSeek in JSON mode and parse the result.

    JSON mode should make fences impossible, but strip them defensively so a
    stray wrapper degrades to a successful parse rather than a 500.
    """
    text = _deepseek_chat(prompt, json_mode=True, max_tokens=max_tokens)

    if text.startswith("```"):
        parts = text.split("```")
        if len(parts) > 1:
            text = parts[1]
        if text.startswith("json"):
            text = text[4:]

    try:
        return json.loads(text.strip())
    except json.JSONDecodeError as e:
        raise DeepSeekError(f"DeepSeek returned malformed JSON: {e}")

# ── ElevenLabs ───────────────────────────────────────────────────────────────

ELEVENLABS_API_KEY = os.environ.get("ELEVENLABS_API_KEY", "")
elevenlabs_client = ElevenLabs(api_key=ELEVENLABS_API_KEY) if ELEVENLABS_API_KEY else None
ELEVENLABS_TTS_URL = "https://api.elevenlabs.io/v1/text-to-speech/{voice_id}/stream"
ELEVENLABS_VOICES_URL = "https://api.elevenlabs.io/v1/voices"

# ── Google OAuth (Login) ────────────────────────────────────────────────────

GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")
REDIRECT_URI = os.environ.get("REDIRECT_URI", "http://localhost:5001/callback")

if not IS_PRODUCTION:
    os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"

CLIENT_SECRETS = {
    "web": {
        "client_id": GOOGLE_CLIENT_ID,
        "client_secret": GOOGLE_CLIENT_SECRET,
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
        "redirect_uris": [REDIRECT_URI],
    }
}

LOCAL_MODE = os.environ.get("LOCAL_MODE", "").lower() == "true"
LOCAL_USER = {"name": "Shah", "email": "aalikes@gmail.com", "picture": ""}

# ── Notion Integration ──────────────────────────────────────────────────────

NOTION_TOKEN           = os.environ.get("NOTION_TOKEN", "")
NOTION_PARENT_ID       = os.environ.get("NOTION_PARENT_PAGE_ID", "")
NOTION_SPEECHES_DB_ID  = "c9dc5a4f20fc4ee2af9968419688dbd7"
NOTION_BRIDGE_URL      = os.environ.get("NOTION_BRIDGE_URL", "")

# ── Google Drive (Export) ────────────────────────────────────────────────────

_app_port         = os.environ.get("PORT", "5001")
DRIVE_TOKEN_FILE  = os.path.join(os.path.dirname(__file__), "drive_token.json")
DRIVE_REDIRECT    = f"{'https' if IS_PRODUCTION else 'http'}://localhost:{_app_port}/auth/drive/callback"
DRIVE_SCOPES      = [
    "https://www.googleapis.com/auth/drive.file",
    "https://www.googleapis.com/auth/drive.readonly",
]
DRIVE_CLIENT_SECRETS = {
    "web": {
        "client_id": os.environ.get("GOOGLE_DRIVE_CLIENT_ID", GOOGLE_CLIENT_ID),
        "client_secret": os.environ.get("GOOGLE_DRIVE_CLIENT_SECRET", GOOGLE_CLIENT_SECRET),
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
        "redirect_uris": [DRIVE_REDIRECT],
    }
}
TOASTMASTERS_FOLDER_NAME = "Toastmasters Workspaces"

# ── Google Apps Script (Google Doc export) ──────────────────────────────────

APPS_SCRIPT_URL = os.environ.get("APPS_SCRIPT_URL", "")

# ── Rate Limiter Helper ─────────────────────────────────────────────────────

from functools import wraps
from flask import current_app

def rate_limit(max_per_minute=60):
    """Simple in-memory rate limiter decorator (no external deps).
    For production, swap to flask-limiter with Redis backend."""
    from collections import defaultdict
    import time as _time

    limiter_buckets = defaultdict(list)

    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            key = f"{request.remote_addr}:{request.path}"
            now = _time.time()
            window = 60.0
            bucket = [t for t in limiter_buckets[key] if now - t < window]
            bucket.append(now)
            limiter_buckets[key] = bucket
            if len(bucket) > max_per_minute:
                return jsonify({"error": "Rate limit exceeded. Try again shortly."}), 429
            return f(*args, **kwargs)
        return wrapped
    return decorator

# ── Persistent Speech Database ────────────────────────────────────────────────

SPEECHES_DB  = os.path.join(os.path.dirname(__file__), "speeches.db")
SYNC_QUEUE   = os.path.join(os.path.dirname(__file__), "sync_queue.json")


def get_db():
    conn = sqlite3.connect(SPEECHES_DB)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with get_db() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS speeches (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                title       TEXT NOT NULL,
                pathway     TEXT,
                project     TEXT,
                language    TEXT DEFAULT 'English',
                topic       TEXT,
                duration    TEXT,
                text        TEXT NOT NULL,
                word_count  INTEGER DEFAULT 0,
                created_at  INTEGER NOT NULL,
                updated_at  INTEGER NOT NULL,
                notion_url  TEXT,
                drive_url   TEXT,
                synced_at   INTEGER
            )
        """)
        # Add sync columns to existing tables that predate this migration
        for col in ("notion_url TEXT", "drive_url TEXT", "synced_at INTEGER"):
            try:
                conn.execute(f"ALTER TABLE speeches ADD COLUMN {col}")
            except Exception:
                pass
        conn.commit()


def _enqueue_sync(speech_id, title, pathway, project, duration, topic,
                  objectives, text):
    """Append a speech to the local sync queue for Cowork to pick up."""
    try:
        queue = []
        if os.path.exists(SYNC_QUEUE):
            with open(SYNC_QUEUE) as f:
                queue = json.load(f)
        # Remove any existing entry for this id to avoid duplicates
        queue = [e for e in queue if e.get("id") != speech_id]
        queue.append({
            "id":         speech_id,
            "title":      title,
            "pathway":    pathway,
            "project":    project,
            "duration":   duration,
            "topic":      topic,
            "objectives": objectives,
            "text":       text,
            "queued_at":  int(time.time()),
        })
        with open(SYNC_QUEUE, "w") as f:
            json.dump(queue, f, indent=2)
    except Exception as e:
        app.logger.warning(f"Sync queue write failed: {e}")


init_db()


# ── Error handlers ──────────────────────────────────────────────────────────────

@app.errorhandler(404)
def not_found(_e):
    user = session.get("user")
    return render_template("404.html", user=user), 404


@app.errorhandler(500)
def server_error(_e):
    user = session.get("user")
    return render_template("500.html", user=user), 500


# ── Auth routes ────────────────────────────────────────────────────────────────

@app.route("/login")
@rate_limit(max_per_minute=30)
def login():
    flow = Flow.from_client_config(
        CLIENT_SECRETS,
        scopes=["openid", "https://www.googleapis.com/auth/userinfo.email",
                "https://www.googleapis.com/auth/userinfo.profile"],
        redirect_uri=REDIRECT_URI,
    )
    auth_url, state = flow.authorization_url(prompt="consent")
    session["state"] = state
    return redirect(auth_url)


@app.route("/callback")
def callback():
    flow = Flow.from_client_config(
        CLIENT_SECRETS,
        scopes=["openid", "https://www.googleapis.com/auth/userinfo.email",
                "https://www.googleapis.com/auth/userinfo.profile"],
        redirect_uri=REDIRECT_URI,
        state=session.get("state"),
    )
    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials
    id_info = id_token.verify_oauth2_token(
        credentials.id_token, google_requests.Request(), GOOGLE_CLIENT_ID
    )
    session["user"] = {
        "name": id_info.get("name"),
        "email": id_info.get("email"),
        "picture": id_info.get("picture"),
    }
    return redirect(url_for("builder"))


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))


# ── Main routes ────────────────────────────────────────────────────────────────

@app.before_request
def auto_login_local():
    """In LOCAL_MODE, inject a default user so OAuth is never needed."""
    if LOCAL_MODE and not session.get("user"):
        session["user"] = LOCAL_USER


@app.route("/")
def index():
    user = session.get("user")
    return render_template("index.html", user=user)


@app.route("/builder")
def builder():
    if not session.get("user"):
        return redirect(url_for("login"))
    return render_template("builder.html", user=session["user"])


@app.route("/archive")
def archive():
    if not session.get("user"):
        return redirect(url_for("login"))
    return render_template("archive.html", user=session["user"])


@app.route("/table-topics")
def table_topics():
    if not session.get("user"):
        return redirect(url_for("login"))
    return render_template("table-topics.html", user=session["user"])


@app.route("/settings")
def settings():
    if not session.get("user"):
        return redirect(url_for("login"))
    drive_creds = _drive_credentials()
    drive_email = ""
    if drive_creds:
        try:
            svc = _drive_service()
            about = svc.about().get(fields="user").execute()
            drive_email = about.get("user", {}).get("emailAddress", "")
        except Exception:
            pass
    return render_template("settings.html", user=session["user"],
                           drive_connected=bool(drive_creds),
                           drive_email=drive_email,
                           notion_configured=bool(NOTION_TOKEN),
                           elevenlabs_configured=bool(ELEVENLABS_API_KEY),
                           drive_redirect=DRIVE_REDIRECT)


@app.route("/import")
def import_speech():
    if not session.get("user"):
        return redirect(url_for("login"))
    return render_template("import.html", user=session["user"])


@app.route("/import/file", methods=["POST"])
def import_file():
    """Extract plain text from an uploaded .txt, .docx, or .pdf file."""
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401

    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file provided"}), 400

    filename = f.filename.lower()
    try:
        if filename.endswith(".txt"):
            text = f.read().decode("utf-8", errors="replace")

        elif filename.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(f)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif filename.endswith(".pdf"):
            from pdfminer.high_level import extract_text as pdf_extract
            from io import BytesIO
            text = pdf_extract(BytesIO(f.read()))

        else:
            return jsonify({"error": "Unsupported file type. Use .txt, .docx, or .pdf"}), 400

        return jsonify({"text": text.strip(), "filename": f.filename})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/import/drive-url", methods=["POST"])
def import_drive_url():
    """Fetch a Google Doc (or any public URL) and return its plain text."""
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401

    import re as _re
    data = request.get_json()
    url  = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400

    # Convert Google Docs share/edit links → export-as-txt URL
    gdoc_match = _re.search(
        r"docs\.google\.com/document/d/([A-Za-z0-9_\-]+)", url
    )
    gdrive_match = _re.search(
        r"drive\.google\.com/file/d/([A-Za-z0-9_\-]+)", url
    )

    fetch_url = url  # default: use as-is
    source_label = "URL"

    if gdoc_match:
        doc_id = gdoc_match.group(1)
        fetch_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
        source_label = "Google Doc"
    elif gdrive_match:
        file_id = gdrive_match.group(1)
        fetch_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        source_label = "Google Drive file"

    try:
        resp = requests.get(fetch_url, timeout=15, headers={"User-Agent": "VoxArtisan/1.0"})
        if resp.status_code == 401:
            return jsonify({"error": (
                "This document requires sign-in. Make sure sharing is set to "
                "'Anyone with the link can view' in Google Docs."
            )}), 403
        if resp.status_code != 200:
            return jsonify({"error": f"Could not fetch {source_label} (HTTP {resp.status_code})"}), 400

        content_type = resp.headers.get("Content-Type", "")
        text = resp.text if "text" in content_type else resp.content.decode("utf-8", errors="replace")
        return jsonify({"text": text.strip(), "source": source_label})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/generate", methods=["POST"])
@rate_limit(max_per_minute=10)
def generate():
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401

    data = request.get_json()
    topic               = data.get("topic", "")
    pathway             = data.get("pathway", "")
    project             = data.get("project", "")
    story_line          = data.get("story_line", "")
    central_message     = data.get("central_message", "")
    intro_style         = data.get("intro_style", "Bold Statement")
    closing_technique   = data.get("closing_technique", "Callback to Opening")
    project_objectives  = data.get("project_objectives", "")
    audience            = data.get("audience", "general")
    duration            = data.get("duration", "5-7 minutes")
    tone                = data.get("tone", "professional")
    persona             = data.get("persona", "")
    key_points          = data.get("key_points", "")
    language            = data.get("language", "English")

    excluded_sections   = data.get("excluded_sections", [])
    custom_additions    = data.get("custom_additions", "")

    clean_pathway       = pathway if pathway and pathway != "— Select a Pathway —" else ""
    pathway_line        = f"- Toastmasters Pathway: {clean_pathway}" if clean_pathway else ""
    project_line        = f"- Project: {project}" if project else ""

    # Extract level from project string e.g. "Level 2 — Effective Body Language" → "Level 2"
    import re as _re
    level_match         = _re.match(r"(Level\s+\d+)", project or "", _re.IGNORECASE)
    level_str           = level_match.group(1) if level_match else ""
    project_name_only   = _re.sub(r"^Level\s+\d+\s*[—\-–]\s*", "", project or "").strip() if project else ""
    story_line_text     = f"\nNarrative / Story Line to weave in:\n{story_line}" if story_line else ""
    central_msg_line    = f"- Central Message & CTA: {central_message}" if central_message else ""
    closing_line        = f"- Closing Technique: {closing_technique}"
    objectives_text     = f"\nProject Objectives to satisfy:\n{project_objectives}" if project_objectives else ""
    key_points_line     = f"- Key Points: {key_points}" if key_points else ""
    persona_note        = _persona_instruction(persona)
    language_note       = _language_instructions(language)
    excluded_note       = f"\n⛔ SKIP THESE SECTIONS ENTIRELY — do not generate them: {', '.join(excluded_sections)}" if excluded_sections else ""
    custom_note         = f"\n🔧 ADDITIONAL INSTRUCTIONS FOR THIS ITERATION:\n{custom_additions}" if custom_additions else ""

    prompt = f"""Act as a World Class Public Speaking Coach and Speechwriter.
Your goal is to write a speech for a Toastmaster using the following constraints:

- Pathway: {clean_pathway or "Not specified"}
- Tone: {tone}
- Style: {intro_style}
- Persona: {persona if persona else "None (default)"}{persona_note}

Speech Brief:
{pathway_line}
{project_line}
- Subject Matter: {topic}
- Audience: {audience}
- Duration: {duration}
- Tone: {tone}
- Opening Technique: {intro_style}
{central_msg_line}
{closing_line}
{key_points_line}
- Language: {language}
{story_line_text}
{objectives_text}
{language_note}
{excluded_note}
{custom_note}

STRICT FORMATTING RULES:
1. Use [brackets] for stage directions (e.g., [Pause for laughter], [Walk to the left stage]).
2. Include a specialized 'Introducer's Introduction' at the very beginning.
3. The 'Hook' must align with the chosen style: {intro_style}.
4. Ensure the 'Call to Action' is powerful and ties back to the 'Central Message'.
5. Vocabulary should be conversational yet professional, avoiding AI-isms like 'delve' or 'tapestry'.

MANDATORY SPEECH STRUCTURE — deliver the sections in this exact order:

## TITLE
A punchy, memorable title for the speech. Generate this FIRST — it will be referenced in the Introducer's Introduction below.

## INTRODUCER'S INTRODUCTION
A card written for the Toastmaster or evaluator who will introduce Shah to the room. MANDATORY — you MUST include ALL FIVE of the following lines verbatim at the top of this section, filled in with the actual values:

  Pathway: {clean_pathway or "[Not specified]"}
  Level: {level_str or "[Not specified]"}
  Project: {project_name_only or "[Not specified]"}
  Time: {duration or "[Not specified]"}
  Title: [insert the exact title you generated above]

After those five lines, write 2–3 sentences in third person ("Our next speaker is Shah..."). Be intentionally vague about the speech's content — build intrigue without revealing the theme or argument. End with exactly: "The title of his speech is '[title]'. Please welcome, Shah!" This section is read by the introducer, not by Shah.

## HOOK
The VERY FIRST WORDS Shah delivers on stage — no salutation before this. Use the "{intro_style}" technique to grab the room in the first 30 seconds.

## SALUTATION
After the hook lands, Shah delivers the Toastmasters salutation:
"Mr. and Mrs. Toastmaster, distinguished Toastmasters, fellow Toastmasters, and guests..."

## BODY
Structured content with smooth transitions that build toward the central message. Include [speaker notes in brackets] for emphasis, pauses, gestures, and vocal variety.

## CLOSE
Use the "{closing_technique}" technique — memorable, emotionally resonant, with a clear call-to-action.

## CLOSING SALUTATION
End with the formal Toastmasters handoff: "Back to you, Mister — or Madam — Toastmaster."

## WORD COUNT
Estimated word count and speaking time."""

    try:
        return jsonify({"speech": _deepseek_chat(prompt)})
    except DeepSeekError as e:
        return jsonify({"error": str(e)}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/suggest", methods=["POST"])
def suggest():
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401

    data = request.get_json()
    topic               = data.get("topic", "")
    pathway             = data.get("pathway", "")
    project             = data.get("project", "")
    tone                = data.get("tone", "")
    story_line          = data.get("story_line", "")
    central_message     = data.get("central_message", "")
    intro_style         = data.get("intro_style", "Bold Statement")
    closing_technique   = data.get("closing_technique", "Callback to Opening")
    duration            = data.get("duration", "5-7 minutes")
    language            = data.get("language", "English")

    context = f"Topic: {topic} | Pathway: {pathway} | Project: {project} | Tone: {tone} | Duration: {duration} | Preferred Intro Style: {intro_style} | Closing Technique: {closing_technique} | Language: {language}"
    if central_message:
        context += f" | Central Message: {central_message}"
    if story_line:
        context += f" | Story: {story_line}"

    lang_instruction = f" Generate ALL titles and intros in {language}." if language != "English" else ""
    if language.lower() in ("haitian creole", "kreyòl", "kreyol"):
        lang_instruction += " Use 'Kreyòl swa' (elegant formal Creole). Avoid literal French-to-Creole translations. Integrate a relevant 'Pwoveb' (proverb) if natural. Ensure warmth in the salutation."

    prompt = f"""You are VoxArtisan, an elite Toastmasters speech coach.

Based on this speech brief:
{context}

Return ONLY valid JSON (no markdown, no explanation) in this exact structure:
{{
  "titles": [
    "Title Option 1",
    "Title Option 2",
    "Title Option 3",
    "Title Option 4",
    "Title Option 5"
  ],
  "intros": [
    "Full opening using the {intro_style} technique — primary recommendation.",
    "Full opening using a different technique as an alternative — complete, speakable.",
    "Full opening using a third technique — bold, unexpected, or story-driven."
  ]
}}

Titles: punchy, memorable, fit the tone and subject.
Intros: complete speakable sentences (2-3 sentences max), first intro MUST use the "{intro_style}" technique, the other two use varied alternatives. IMPORTANT: These are HOOK openings only — no Toastmasters salutation yet.{lang_instruction}"""

    try:
        return jsonify(_deepseek_json(prompt))
    except DeepSeekError as e:
        return jsonify({"error": str(e)}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/proofread", methods=["POST"])
def proofread():
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401

    data   = request.get_json()
    speech = data.get("speech", "")

    prompt = f"""You are an expert Toastmasters speech coach and editor. Carefully proofread the speech below.

Return ONLY valid JSON — no markdown, no explanation — in exactly this structure:
{{
  "score": 82,
  "overall": "One or two sentence assessment of the speech's current state.",
  "strengths": ["Specific strength 1", "Specific strength 2", "Specific strength 3"],
  "suggestions": [
    {{
      "type": "grammar",
      "original": "exact phrase from the speech to replace (keep short, max 12 words)",
      "suggestion": "improved replacement text",
      "reason": "brief explanation (max 10 words)"
    }}
  ],
  "revised": "The complete speech with ALL suggestions applied. Must be the full speech text."
}}

Types: "grammar", "clarity", "word_choice", "pacing", "toastmasters"
Score: 0–100 readiness score (100 = stage-ready).
Include 4–8 of the most impactful suggestions only — no nitpicking.
The "original" field must be an EXACT verbatim substring from the speech (used for find-and-replace).

SPEECH:
{speech}"""

    try:
        return jsonify(_deepseek_json(prompt, max_tokens=16384))
    except DeepSeekError as e:
        return jsonify({"error": str(e)}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/refine", methods=["POST"])
def refine():
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401

    data = request.get_json()
    original = data.get("speech", "")
    instruction = data.get("instruction", "")

    prompt = f"""You are an expert Toastmasters coach. Refine this speech based on the instruction below.

ORIGINAL SPEECH:
{original}

REFINEMENT INSTRUCTION:
{instruction}

Return the improved full speech text only."""

    try:
        return jsonify({"speech": _deepseek_chat(prompt)})
    except DeepSeekError as e:
        return jsonify({"error": str(e)}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Persona-specific instructions ────────────────────────────────────────────

def _persona_instruction(persona):
    """Return storytelling persona instructions for the speech generation."""
    prompts = {
        "Reynolds": "\n\nINFUSE the speech with the persona of Ryan Reynolds (Deadpool): Use a sarcastic, fast-paced delivery. Add self-deprecating asides in [brackets]. Keep the tone irreverent but charming. The humor should land like a witty action-comedy.",
        "Hart": "\n\nINFUSE the speech with the persona of Kevin Hart: High-energy and frantic relatability. Use repetition and exclamation for comedic timing. Make it feel like a story a friend is telling you at a party — loud, animated, and deeply personal.",
        "Chappelle": "\n\nINFUSE the speech with the persona of Dave Chappelle: Reflective and poignant. Focus on thematic weight. Use masterful pauses [Long Pause] and a cool, observational storytelling style. Let silences breathe. The humor comes from truth, not punchlines.",
    }
    return prompts.get(persona, "")


# ── Language-specific instructions ─────────────────────────────────────────────

def _language_instructions(language):
    """Return language-specific prompt instructions for the speech generation."""
    base = f"\n⚠️ LANGUAGE REQUIREMENT: Generate the ENTIRE speech — every word, every speaker note — in {language}. Do not use English."
    if language.lower() in ("haitian creole", "kreyòl", "kreyol"):
        return base + """

If the language selected is Haitian Creole (Kreyòl):
- DO NOT use literal French-to-Creole translations.
- USE 'Kreyòl swa' (elegant, formal Creole) for the body of the speech.
- INTEGRATE at least one relevant 'Pwoveb' (proverb) that aligns with the central message.
- TONE: Ensure the 'Salutation' reflects the warmth of a Haitian community gathering.
- IDIOMS: Use natural expressions like 'voye pwen' or 'bat bouch' appropriately if the tone is 'Humorous'."""
    return base


# ── Google Drive OAuth ────────────────────────────────────────────────────────

def _drive_credentials():
    """Load saved Drive credentials, refreshing if expired. Returns None if not connected."""
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request as GRequest
    if not os.path.exists(DRIVE_TOKEN_FILE):
        return None
    try:
        creds = Credentials.from_authorized_user_file(DRIVE_TOKEN_FILE, DRIVE_SCOPES)
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(GRequest())
            with open(DRIVE_TOKEN_FILE, "w") as f:
                f.write(creds.to_json())
        return creds if creds and creds.valid else None
    except Exception:
        return None


def _drive_service():
    """Return an authenticated Drive service, or None."""
    creds = _drive_credentials()
    if not creds:
        return None
    return build("drive", "v3", credentials=creds)


def _get_or_create_folder(service, name, parent_id=None):
    """Return the ID of a Drive folder, creating it if it doesn't exist."""
    query = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    results = service.files().list(q=query, fields="files(id,name)", pageSize=1).execute()
    files = results.get("files", [])
    if files:
        return files[0]["id"]
    meta = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
    if parent_id:
        meta["parents"] = [parent_id]
    folder = service.files().create(body=meta, fields="id").execute()
    return folder["id"]


@app.route("/auth/drive")
def auth_drive():
    if not session.get("user"):
        return redirect(url_for("login"))
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        return redirect(url_for("settings") + "?error=missing_credentials")
    flow = Flow.from_client_config(
        DRIVE_CLIENT_SECRETS,
        scopes=DRIVE_SCOPES,
        redirect_uri=DRIVE_REDIRECT,
    )
    auth_url, state = flow.authorization_url(
        prompt="consent", access_type="offline", include_granted_scopes="true"
    )
    session["drive_state"] = state
    return redirect(auth_url)


@app.route("/auth/drive/callback")
def auth_drive_callback():
    state = session.get("drive_state")
    flow = Flow.from_client_config(
        DRIVE_CLIENT_SECRETS,
        scopes=DRIVE_SCOPES,
        redirect_uri=DRIVE_REDIRECT,
        state=state,
    )
    flow.fetch_token(authorization_response=request.url)
    creds = flow.credentials
    with open(DRIVE_TOKEN_FILE, "w") as f:
        f.write(creds.to_json())
    return redirect(url_for("settings") + "?connected=drive")


@app.route("/auth/drive/disconnect", methods=["POST"])
def auth_drive_disconnect():
    if os.path.exists(DRIVE_TOKEN_FILE):
        os.remove(DRIVE_TOKEN_FILE)
    return redirect(url_for("settings") + "?disconnected=drive")


@app.route("/drive/status")
def drive_status():
    creds = _drive_credentials()
    connected = creds is not None
    email = ""
    if connected:
        try:
            svc = _drive_service()
            about = svc.about().get(fields="user").execute()
            email = about.get("user", {}).get("emailAddress", "")
        except Exception:
            pass
    return jsonify({"connected": connected, "email": email})


@app.route("/drive/files")
def drive_files():
    """List speech files from the Toastmasters Workspaces folder."""
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401
    svc = _drive_service()
    if not svc:
        return jsonify({"error": "drive_not_connected"}), 403
    try:
        folder_id = _get_or_create_folder(svc, TOASTMASTERS_FOLDER_NAME)
        query = (
            f"'{folder_id}' in parents and trashed=false and ("
            "mimeType='application/vnd.google-apps.document' or "
            "mimeType='text/plain' or "
            "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'"
            ")"
        )
        results = svc.files().list(
            q=query,
            fields="files(id,name,mimeType,modifiedTime,size)",
            orderBy="modifiedTime desc",
            pageSize=50,
        ).execute()
        return jsonify({"files": results.get("files", []), "folder_id": folder_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/drive/download/<file_id>")
def drive_download(file_id):
    """Fetch text content of a Drive file for import."""
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401
    svc = _drive_service()
    if not svc:
        return jsonify({"error": "drive_not_connected"}), 403
    try:
        meta = svc.files().get(fileId=file_id, fields="name,mimeType").execute()
        mime = meta.get("mimeType", "")
        name = meta.get("name", "")
        # Export Google Docs as plain text
        if mime == "application/vnd.google-apps.document":
            content = svc.files().export(fileId=file_id, mimeType="text/plain").execute()
            text = content.decode("utf-8") if isinstance(content, bytes) else content
        else:
            import io as _io
            from googleapiclient.http import MediaIoBaseDownload
            fh = _io.BytesIO()
            dl = MediaIoBaseDownload(fh, svc.files().get_media(fileId=file_id))
            done = False
            while not done:
                _, done = dl.next_chunk()
            raw = fh.getvalue()
            if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                from docx import Document as DocxDoc
                import io as _io2
                doc = DocxDoc(_io2.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                text = raw.decode("utf-8", errors="replace")
        return jsonify({"text": text.strip(), "name": name})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/drive/upload/<int:speech_id>", methods=["POST"])
def drive_upload(speech_id):
    """Upload a saved speech to the Toastmasters Workspaces folder."""
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401
    svc = _drive_service()
    if not svc:
        return jsonify({"error": "drive_not_connected"}), 403
    with get_db() as conn:
        row = conn.execute("SELECT * FROM speeches WHERE id=?", (speech_id,)).fetchone()
    if not row:
        return jsonify({"error": "Speech not found"}), 404
    speech = dict(row)
    try:
        folder_id = _get_or_create_folder(svc, TOASTMASTERS_FOLDER_NAME)
        content = speech["text"].encode("utf-8")
        meta = {
            "name": speech["title"],
            "parents": [folder_id],
            "mimeType": "application/vnd.google-apps.document",
        }
        media = MediaIoBaseUpload(BytesIO(content), mimetype="text/plain", resumable=True)
        # Check if already uploaded (drive_url stored)
        if speech.get("drive_url"):
            # Update existing file
            existing_id = speech["drive_url"].split("/d/")[-1].split("/")[0]
            try:
                updated = svc.files().update(
                    fileId=existing_id,
                    body={"name": speech["title"]},
                    media_body=media,
                    fields="id,webViewLink",
                ).execute()
                url = updated.get("webViewLink", "")
            except Exception:
                # File may have been deleted — create new
                f = svc.files().create(body=meta, media_body=media, fields="id,webViewLink").execute()
                url = f.get("webViewLink", "")
        else:
            f = svc.files().create(body=meta, media_body=media, fields="id,webViewLink").execute()
            url = f.get("webViewLink", "")
        # Persist drive_url
        with get_db() as conn:
            conn.execute("UPDATE speeches SET drive_url=?, synced_at=? WHERE id=?",
                         (url, int(time.time()), speech_id))
            conn.commit()
        return jsonify({"url": url})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Export: Notion ────────────────────────────────────────────────────────────

def _build_notion_blocks(speech_text):
    """Convert speech text into Notion block objects."""
    blocks = []
    for line in speech_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## ") or line.startswith("# "):
            heading = line.lstrip("# ").strip()
            blocks.append({
                "object": "block", "type": "heading_2",
                "heading_2": {"rich_text": [{"type": "text", "text": {"content": heading[:2000]}}]}
            })
        elif line.startswith("[") and line.endswith("]"):
            blocks.append({
                "object": "block", "type": "callout",
                "callout": {
                    "rich_text": [{"type": "text", "text": {"content": line[1:-1][:2000]}}],
                    "icon": {"emoji": "🎤"}
                }
            })
        else:
            blocks.append({
                "object": "block", "type": "paragraph",
                "paragraph": {"rich_text": [{"type": "text", "text": {"content": line[:2000]}}]}
            })
    return blocks


def _extract_section(text, section_name):
    """Extract a named ## SECTION block from speech text."""
    import re
    pattern = rf"##\s*{re.escape(section_name)}\s*\n(.*?)(?=\n##\s|\Z)"
    m = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    return m.group(1).strip() if m else ""


def _map_duration(duration_str):
    """Map free-text duration to Notion select option."""
    mapping = {
        "4": "4-6 mins", "5": "5-7 mins",
        "8": "8-10 mins", "10": "10-12 mins", "13": "13-15 mins",
    }
    for key, val in mapping.items():
        if key in str(duration_str):
            return val
    return None


@app.route("/export/notion", methods=["POST"])
def export_notion():
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401
    if not NOTION_TOKEN:
        return jsonify({"error": "notion_not_configured"}), 400

    data     = request.get_json()
    speech   = data.get("speech", "")
    title    = data.get("title", "VoxArtisan Speech").strip()
    pathway  = data.get("pathway", "")
    project  = data.get("project", "")
    duration = data.get("duration", "")
    topic    = data.get("topic", "")
    objectives = data.get("project_objectives", "")

    # Extract structured sections
    intro_text = _extract_section(speech, "INTRODUCER'S INTRODUCTION")
    hook_text  = _extract_section(speech, "HOOK")

    # Build properties for the Speeches database
    properties = {
        "Speech Title": {
            "title": [{"text": {"content": title[:200]}}]
        },
        "Status": {
            "select": {"name": "In Progress"}
        },
        "Assigned To": {
            "rich_text": [{"text": {"content": "Shah"}}]
        },
    }

    if intro_text:
        properties["Speech Intro"] = {
            "rich_text": [{"text": {"content": intro_text[:2000]}}]
        }
    if hook_text:
        ai_summary = hook_text[:500]
        properties["AI Summary"] = {
            "rich_text": [{"text": {"content": ai_summary}}]
        }
    if objectives:
        properties["Objectives"] = {
            "rich_text": [{"text": {"content": objectives[:2000]}}]
        }
    if topic:
        properties["Notes"] = {
            "rich_text": [{"text": {"content": f"Topic: {topic}"}}]
        }

    duration_val = _map_duration(duration)
    if duration_val:
        properties["Speech Duration"] = {"select": {"name": duration_val}}

    # Build page body blocks
    blocks = _build_notion_blocks(speech)

    try:
        notion = NotionClient(auth=NOTION_TOKEN)
        # Create entry directly in the Speeches database
        page = notion.pages.create(
            parent={"database_id": NOTION_SPEECHES_DB_ID},
            properties=properties,
            children=blocks[:100],
        )
        page_url = page.get("url", "")
        page_id  = page.get("id", "")

        # If there are more than 100 blocks, append the rest
        if len(blocks) > 100:
            for i in range(100, len(blocks), 100):
                notion.blocks.children.append(
                    block_id=page_id,
                    children=blocks[i:i+100]
                )

        return jsonify({"url": page_url, "id": page_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/export/gdoc", methods=["POST"])
def export_gdoc():
    """POST speech data to the Google Apps Script Web App, which creates a formatted Google Doc and returns the URL."""
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401
    if not APPS_SCRIPT_URL:
        return jsonify({"error": "APPS_SCRIPT_URL not configured. Set it in .env"}), 400

    data = request.get_json()
    if not data or not data.get("content"):
        return jsonify({"error": "No speech content provided"}), 400

    # Extract the introducer section from the full speech
    intro = _extract_section(data.get("content", ""), "INTRODUCER'S INTRODUCTION")
    # Strip the introducer block from the content to get the body
    body = data.get("content", "")
    if intro:
        # Remove the ## INTRODUCER'S INTRODUCTION block
        import re as _re
        body = _re.sub(
            r"^##\s*INTRODUCER'S INTRODUCTION\s*\n[\s\S]*?(?=\n##\s|$)",
            "",
            body,
            flags=_re.MULTILINE | _re.IGNORECASE,
        ).strip()

    payload = {
        "title": data.get("title", "VoxArtisan Speech"),
        "pathway": data.get("pathway", ""),
        "project": data.get("project", ""),
        "duration": data.get("duration", ""),
        "tone": data.get("tone", ""),
        "topic": data.get("topic", ""),
        "introducer": intro or "",
        "content": body or data.get("content", ""),
    }

    try:
        import requests as _req
        resp = _req.post(
            APPS_SCRIPT_URL,
            json=payload,
            headers={"Content-Type": "application/json"},
            timeout=30,
        )
        result = resp.json()
        if resp.ok and result.get("status") == "success":
            return jsonify({"url": result["url"]})
        else:
            return jsonify({"error": result.get("error", "Apps Script request failed")}), 502
    except Exception as e:
        return jsonify({"error": f"Google Doc export failed: {str(e)}"}), 500


# ── ElevenLabs TTS ───────────────────────────────────────────────────────────

@app.route("/tts", methods=["POST"])
def tts_generate():
    if not session.get("user"):
        return jsonify({"error": "Not authenticated"}), 401
    if not ELEVENLABS_API_KEY:
        return jsonify({"error": "elevenlabs_not_configured"}), 400

    data       = request.get_json()
    text       = data.get("text", "")[:12000]   # hard cap ~90-min speech
    voice_id   = data.get("voice_id", "21m00Tcm4TlvDq8ikWAM")
    stability  = float(data.get("stability", 0.45))
    similarity = float(data.get("similarity", 0.80))

    url = ELEVENLABS_TTS_URL.format(voice_id=voice_id)
    headers = {
        "xi-api-key": ELEVENLABS_API_KEY,
        "Content-Type": "application/json",
        "Accept": "audio/mpeg",
    }
    payload = {
        "text": text,
        "model_id": "eleven_multilingual_v2",
        "voice_settings": {
            "stability": stability,
            "similarity_boost": similarity,
            "style": 0.0,
            "use_speaker_boost": True,
        },
    }

    try:
        resp = requests.post(url, json=payload, headers=headers, stream=True, timeout=60)
        if resp.status_code != 200:
            msg = resp.text[:300]
            return jsonify({"error": f"ElevenLabs {resp.status_code}: {msg}"}), 500

        def generate():
            for chunk in resp.iter_content(chunk_size=4096):
                if chunk:
                    yield chunk

        return Response(
            stream_with_context(generate()),
            content_type="audio/mpeg",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/tts/voices")
def tts_voices():
    """Return available ElevenLabs voices for the account."""
    if not ELEVENLABS_API_KEY:
        return jsonify({"error": "elevenlabs_not_configured"}), 400
    try:
        resp = requests.get(
            ELEVENLABS_VOICES_URL,
            headers={"xi-api-key": ELEVENLABS_API_KEY},
            timeout=10,
        )
        voices = resp.json().get("voices", [])
        return jsonify([
            {"id": v["voice_id"], "name": v["name"],
             "labels": v.get("labels", {}), "preview_url": v.get("preview_url", "")}
            for v in voices
        ])
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Saved Speeches CRUD ───────────────────────────────────────────────────────

@app.route("/speeches", methods=["GET"])
def list_speeches():
    """Return all saved speeches, newest first."""
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, title, pathway, project, language, topic, duration, word_count, created_at, updated_at "
            "FROM speeches ORDER BY updated_at DESC"
        ).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/speeches", methods=["POST"])
def save_speech():
    """Save a new speech or overwrite an existing one by id."""
    data       = request.get_json()
    speech_id  = data.get("id")          # if provided → update
    title      = (data.get("title") or "Untitled Speech").strip()
    pathway    = data.get("pathway", "")
    project    = data.get("project", "")
    language   = data.get("language", "English")
    topic      = data.get("topic", "")
    duration   = data.get("duration", "")
    text       = data.get("text", "")
    word_count = len(text.split())
    now        = int(time.time())

    objectives = data.get("project_objectives", "")

    with get_db() as conn:
        if speech_id:
            conn.execute(
                "UPDATE speeches SET title=?, pathway=?, project=?, language=?, topic=?, duration=?, "
                "text=?, word_count=?, updated_at=?, synced_at=NULL WHERE id=?",
                (title, pathway, project, language, topic, duration, text, word_count, now, speech_id)
            )
            sid = speech_id
        else:
            cur = conn.execute(
                "INSERT INTO speeches (title, pathway, project, language, topic, duration, text, word_count, created_at, updated_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (title, pathway, project, language, topic, duration, text, word_count, now, now)
            )
            sid = cur.lastrowid
        conn.commit()

    # Queue for background Notion + Drive sync
    _enqueue_sync(sid, title, pathway, project, duration, topic, objectives, text)

    return jsonify({"id": sid, "title": title, "word_count": word_count}), 200


@app.route("/speeches/<int:speech_id>/mark_synced", methods=["POST"])
def mark_synced(speech_id):
    """Called by the Cowork sync task after pushing to Notion/Drive."""
    data       = request.get_json()
    notion_url = data.get("notion_url", "")
    drive_url  = data.get("drive_url", "")
    now        = int(time.time())
    with get_db() as conn:
        conn.execute(
            "UPDATE speeches SET notion_url=?, drive_url=?, synced_at=? WHERE id=?",
            (notion_url, drive_url, now, speech_id)
        )
        conn.commit()
    return jsonify({"ok": True})


@app.route("/speeches/<int:speech_id>", methods=["GET"])
def get_speech(speech_id):
    """Fetch a single saved speech by id."""
    with get_db() as conn:
        row = conn.execute("SELECT * FROM speeches WHERE id=?", (speech_id,)).fetchone()
    if not row:
        return jsonify({"error": "not found"}), 404
    return jsonify(dict(row))


@app.route("/sync/queue", methods=["GET"])
def get_sync_queue():
    """Return pending sync queue entries (for Cowork scheduled task)."""
    if not os.path.exists(SYNC_QUEUE):
        return jsonify([])
    with open(SYNC_QUEUE) as f:
        return jsonify(json.load(f))


@app.route("/sync/queue/clear", methods=["POST"])
def clear_sync_queue():
    """Remove synced entries from the queue."""
    data = request.get_json()
    completed_ids = set(data.get("ids", []))
    queue = []
    if os.path.exists(SYNC_QUEUE):
        with open(SYNC_QUEUE) as f:
            queue = json.load(f)
    remaining = [e for e in queue if e.get("id") not in completed_ids]
    with open(SYNC_QUEUE, "w") as f:
        json.dump(remaining, f, indent=2)
    return jsonify({"cleared": len(completed_ids), "remaining": len(remaining)})


@app.route("/speeches/<int:speech_id>", methods=["DELETE"])
def delete_speech(speech_id):
    """Delete a saved speech."""
    with get_db() as conn:
        conn.execute("DELETE FROM speeches WHERE id=?", (speech_id,))
        conn.commit()
    return jsonify({"ok": True})


# ── Notion Worklog tether ─────────────────────────────────────────────────────

VOXARTISAN_WORKLOG_DB_ID = "2dd29dcd-0b81-4b43-aba7-8716dec748a1"


@app.route("/notion/worklog", methods=["POST"])
def notion_worklog():
    """Log a VoxArtisan app event to the Improvements & Worklog Notion database."""
    if not NOTION_TOKEN:
        return jsonify({"ok": False, "error": "notion_not_configured"}), 400

    data   = request.get_json()
    item   = (data.get("item") or "Speech saved").strip()
    notes  = data.get("notes", "")
    area   = data.get("area", "Product")    # "Product","UX","Backend","Frontend","AI/Voice","Ops"
    kind   = data.get("kind", "Implementation")  # "Improvement","Implementation","Iteration","Bug","Decision"
    status = data.get("status", "Done")

    from datetime import date as _date

    try:
        notion = NotionClient(auth=NOTION_TOKEN)
        page = notion.pages.create(
            parent={"database_id": VOXARTISAN_WORKLOG_DB_ID},
            properties={
                "Item": {
                    "title": [{"text": {"content": item}}]
                },
                "Kind": {
                    "select": {"name": kind}
                },
                "Area": {
                    "multi_select": [{"name": area}]
                },
                "Status": {
                    "status": {"name": status}
                },
                "Date": {
                    "date": {"start": _date.today().isoformat()}
                },
                "Notes": {
                    "rich_text": [{"text": {"content": notes}}] if notes else []
                },
            },
        )
        return jsonify({"ok": True, "url": page.get("url", "")})
    except Exception as e:
        app.logger.warning(f"Notion worklog write failed: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


# ── Status checks ─────────────────────────────────────────────────────────────

@app.route("/api/health")
def api_health():
    """Simple health check."""
    return jsonify({"status": "ok", "timestamp": time.time()})


@app.route("/status/integrations")
def integration_status():
    """Return which integrations are configured/connected."""
    return jsonify({
        "gdrive":       _drive_credentials() is not None,
        "notion":       bool(NOTION_TOKEN),
        "elevenlabs":   bool(ELEVENLABS_API_KEY),
        "gscript_url":  APPS_SCRIPT_URL,
        "notion_bridge": NOTION_BRIDGE_URL,
    })


# ── API Dashboard ────────────────────────────────────────────────────────────────

@app.route("/api-dashboard")
def api_dashboard():
    """Render the API monitoring dashboard page."""
    if not session.get("user"):
        return redirect(url_for("login"))
    return render_template("api-dashboard.html", user=session["user"])


# ── Through Line Refinery ────────────────────────────────────────────────────────

@app.route("/through-line")
def through_line():
    """Render the Through Line Refinery page."""
    if not session.get("user"):
        return redirect(url_for("login"))
    return render_template("through-line.html", user=session["user"])


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
